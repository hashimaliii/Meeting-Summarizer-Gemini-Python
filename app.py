import streamlit as st
import tempfile
import os
import io
import math
from datetime import datetime
from docx import Document
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from pydub import AudioSegment
from google import genai

# 1. Setup and Date Extraction
load_dotenv()
client = genai.Client()

# Fixes the [Current Date] placeholder issue
current_date_val = datetime.now().strftime("%B %d, %Y")

# Fetch the system prompt from the .env file
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")

CONCISE_PROMPT = os.getenv("CONCISE_PROMPT", """You are a Chief of Staff. Summarize these notes into a 'Flash Report'. Focus on the 3 most critical outcomes, the 5 most urgent action items, and major blockers. Keep it under 300 words.

MANDATORY DATE: [Insert Today's Date]

FLEXIBLE STRUCTURE RULE:
You may rename or create new headings if they more accurately reflect the urgency or nature of the specific meeting.

MANDATORY EXPORT FORMATTING:
1. CLEAN EXPORT: Use only plain text and simple bullets (-). No tables or bolding inside text blocks.
2. PDF ALIGNMENT: Ensure no line is excessively long; use standard spacing to ensure the PDF export is clean and fully readable without truncation.

Default Structure:
# Executive Flash Report
## I. Key Strategic Outcomes
## II. Urgent Actions
## III. Critical Blockers""")

PROFESSIONAL_PROMPT = os.getenv("PROFESSIONAL_PROMPT", """You are a Senior Corporate Secretary. Your task is to generate exhaustive, professional meeting minutes. Capture the nuance, data points, and all perspectives shared.

MANDATORY DATE: [Insert Today's Date]

FLEXIBLE STRUCTURE RULE: 
If the structure below does not fit the flow of the specific meeting, you are AUTHORIZED and encouraged to generate your own logical headings and sub-headings that best organize the information discussed. 

MANDATORY EXPORT FORMATTING:
1. UNIVERSAL COMPATIBILITY: No Markdown tables, nested columns, or complex symbols.
2. PDF/WORD SAFETY: Use plain text with simple bullet points (-) only. Do NOT use bolding (**) or italics (*) within sentences, as these break the PDF rendering engine and cause text truncation.
3. TEXT WRAPPING: Use clear line breaks and keep paragraphs concise so text wraps correctly in PDF/Word without cutting off at the margins.

Default Structure (Adapt as needed):
# Official Meeting Minutes
## 1. Executive Summary
## 2. Detailed Discussion Points (Create custom sub-headings based on topics)
## 3. Key Decisions Made
## 4. Action Items (Task | Owner | Deadline)
## 5. Next Steps / Follow-up""")

# --- PARSERS FOR MULTIPLE MEDIUMS ---

def extract_text_from_pdf(file_obj):
    reader = PdfReader(file_obj)
    return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

def extract_text_from_docx(file_obj):
    doc = Document(file_obj)
    return "\n".join([para.text for para in doc.paragraphs])

def process_audio(file_obj):
    ext = file_obj.name.split('.')[-1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(file_obj.read())
        path = tmp.name
    try:
        audio = AudioSegment.from_file(path)
        chunk_ms = 45 * 60 * 1000 
        chunks_count = math.ceil(len(audio) / chunk_ms)
        summaries = []
        for i in range(chunks_count):
            chunk = audio[i*chunk_ms : min((i+1)*chunk_ms, len(audio))]
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as c_tmp:
                chunk.export(c_tmp.name, format="mp3")
                g_file = client.files.upload(file=c_tmp.name)
                res = client.models.generate_content(
                    model=GEMINI_MODEL,
                    contents=["Summarize this meeting audio segment in detail.", g_file]
                )
                summaries.append(res.text)
                client.files.delete(name=g_file.name)
        return "\n".join(summaries)
    finally:
        if os.path.exists(path): os.remove(path)

# --- HARDENED EXPORT ENGINES (FIXES TRUNCATION) ---

def create_pdf(text, title="Official Document"):
    """Uses reportlab with proper margins and text wrapping to prevent truncation."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('DocTitle', parent=styles['Title'], fontSize=16, spaceAfter=16)
    heading_style = ParagraphStyle('DocHeading', parent=styles['Heading2'], fontSize=13, spaceBefore=10, spaceAfter=6)
    body_style = ParagraphStyle('DocBody', parent=styles['Normal'], fontSize=11, leading=16, spaceAfter=4)
    bullet_style = ParagraphStyle('DocBullet', parent=styles['Normal'], fontSize=11, leading=16, leftIndent=20, spaceAfter=4)

    story = [Paragraph(title, title_style), Spacer(1, 6)]

    for line in text.split('\n'):
        clean = line.strip().replace('**', '').replace('*', '')
        clean = clean.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        clean_text = clean.lstrip('#').strip()
        if not clean_text:
            story.append(Spacer(1, 6))
            continue
        if clean.startswith('#'):
            story.append(Paragraph(clean_text, heading_style))
        elif clean_text.startswith('-') or clean_text.startswith('\u2022'):
            story.append(Paragraph(clean_text[1:].strip(), bullet_style))
        else:
            story.append(Paragraph(clean_text, body_style))

    doc.build(story)
    return buf.getvalue()

def create_docx(text, title="Meeting Minutes"):
    doc = Document()
    doc.add_heading(title, 0)
    for line in text.split('\n'):
        line = line.strip().replace('**', '').replace('#', '')
        if line:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- UI AND LAYOUT ---
st.set_page_config(page_title="Minutes AI", layout="wide")

import html as _html

def strip_markdown(text):
    """Strip markdown symbols for clean plain-text display."""
    import re
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&family=Barlow:wght@400;500;600;700;800&display=swap');

:root {
    --red:       #E50914;
    --red-dark:  #B20710;
    --bg:        #141414;
    --surface:   #1c1c1c;
    --surface2:  #191919;
    --surface3:  #2a2a2a;
    --border:    rgba(255,255,255,0.07);
    --border2:   rgba(255,255,255,0.13);
    --text:      #e5e5e5;
    --text-dim:  #999;
    --text-muted:#555;
    --white:     #fff;
    --nav-h:     50px;
    --ticker-h:  26px;
    --chrome-h:  76px; /* nav + ticker total */
}

*, *::before, *::after { box-sizing: border-box; }

html, body, [class*="css"] {
    font-family: 'Barlow', sans-serif !important;
    -webkit-font-smoothing: antialiased;
}

/* ── HIDE STREAMLIT CHROME ── */
#MainMenu, footer, header[data-testid="stHeader"],
[data-testid="stToolbar"], [data-testid="stDecoration"],
section[data-testid="stSidebar"] { display: none !important; }

/* ── RESET STREAMLIT WRAPPERS ── */
.stApp { background: var(--bg) !important; color: var(--text) !important; }
.block-container { padding: 0 !important; max-width: 100% !important; }
.stApp > div { padding: 0 !important; }

/* Kill ALL vertical block gaps everywhere */
[data-testid="stVerticalBlock"] { gap: 0 !important; row-gap: 0 !important; }
[data-testid="stVerticalBlock"] > div { padding: 0 !important; margin: 0 !important; }
[data-testid="stVerticalBlockBorderWrapper"] { padding: 0 !important; margin: 0 !important; }

/* ── ANIMATIONS ── */
@keyframes fadeUp  { from { opacity:0; transform:translateY(12px) } to { opacity:1; transform:translateY(0) } }
@keyframes fadeIn  { from { opacity:0 } to { opacity:1 } }
@keyframes ticker  { from { transform:translateX(0) } to { transform:translateX(-50%) } }
@keyframes glowred { 0%,100% { opacity:.4 } 50% { opacity:.9 } }
@keyframes pulsedot { 0%,100% { box-shadow:0 0 0 0 rgba(229,9,20,.55) } 70% { box-shadow:0 0 0 8px rgba(229,9,20,0) } }

/* ── NAVBAR ── */
.nf-nav {
    display: flex; align-items: center; justify-content: space-between;
    padding: 0 1.5rem; height: var(--nav-h);
    background: rgba(12,12,12,0.98);
    border-bottom: 1px solid var(--border);
    position: sticky; top: 0; z-index: 999;
}
.nf-logo {
    font-family: 'Bebas Neue', sans-serif; font-size: 1.5rem;
    letter-spacing: .06em; color: var(--red);
    text-shadow: 0 2px 10px rgba(229,9,20,.3); line-height: 1;
}
.nf-logo span {
    color: rgba(255,255,255,.75); font-size: .65rem;
    letter-spacing: .2em; margin-left: .4rem;
    font-family: 'Barlow', sans-serif; font-weight: 600;
    vertical-align: middle;
}
.nf-nav-right { display:flex; align-items:center; gap:.75rem; }
.nf-date { font-size:.72rem; color:var(--text-dim); font-weight:500; }
.nf-pill {
    font-size:.58rem; font-weight:700; letter-spacing:.1em;
    text-transform:uppercase; color:var(--white); background:var(--red);
    padding:.2rem .55rem; border-radius:2px;
}

/* ── TICKER ── */
.nf-ticker {
    background: rgba(0,0,0,.6); border-bottom: 1px solid var(--border);
    height: var(--ticker-h); overflow: hidden;
    display: flex; align-items: center;
}
.nf-ticker-track { display:flex; width:max-content; animation: ticker 30s linear infinite; }
.nf-ticker-item {
    white-space:nowrap; font-size:.59rem; font-weight:700;
    letter-spacing:.12em; text-transform:uppercase;
    color:var(--text-muted); padding:0 2rem;
}
.nf-ticker-item b { color:var(--red); margin-right:.35rem; }

/* ══════════════════════════════════════════════════
   MAIN TWO-COLUMN LAYOUT
   Streamlit's stHorizontalBlock = our two columns
══════════════════════════════════════════════════ */

/* The outer horizontal block = full-height split layout */
[data-testid="stHorizontalBlock"]:first-of-type {
    height: calc(100vh - var(--chrome-h)) !important;
    align-items: stretch !important;
    gap: 0 !important;
    flex-wrap: nowrap !important;
}
[data-testid="stHorizontalBlock"]:first-of-type > [data-testid="column"] {
    height: 100% !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    padding: 0 !important;
    min-width: 0 !important;
}
/* Left column styling */
[data-testid="stHorizontalBlock"]:first-of-type > [data-testid="column"]:first-child {
    border-right: 1px solid var(--border) !important;
    background: linear-gradient(160deg, #0e0e0e 0%, #141414 100%) !important;
}

/* Inner download-button 3-col row — must NOT stretch to full height */
[data-baseweb="tab-panel"] [data-testid="stHorizontalBlock"] {
    height: auto !important;
    align-items: stretch !important;
    gap: 0 !important;
    padding: 0 !important;
    margin: 0 !important;
}
[data-baseweb="tab-panel"] [data-testid="stHorizontalBlock"] > [data-testid="column"] {
    height: auto !important;
    overflow: visible !important;
    padding: 0 !important;
}

/* ── LEFT PANEL ── */
.left-panel {
    padding: 1.75rem 1.5rem;
    display: flex; flex-direction: column;
    gap: .85rem;
    position: relative; overflow: hidden;
    min-height: 100%;
}
.left-panel::before {
    content:''; position:absolute;
    width:300px; height:300px;
    background:radial-gradient(circle, rgba(229,9,20,.07) 0%, transparent 70%);
    top:-80px; left:-80px; pointer-events:none;
    animation: glowred 5s ease-in-out infinite;
}
.panel-eyebrow {
    font-size:.57rem; font-weight:700; letter-spacing:.18em; text-transform:uppercase;
    color:var(--red); display:flex; align-items:center; gap:.5rem;
}
.panel-eyebrow::before { content:''; width:18px; height:2px; background:var(--red); display:block; flex-shrink:0; }
.panel-title {
    font-family:'Bebas Neue', sans-serif;
    font-size: clamp(1.7rem, 2.8vw, 2.5rem);
    font-weight:400; line-height:.92;
    color:var(--white); letter-spacing:.02em;
}
.panel-title em { color:var(--red); font-style:normal; }
.panel-desc {
    font-size:.79rem; color:var(--text-dim); line-height:1.55; font-weight:400;
}
.fmt-row { display:flex; gap:.3rem; flex-wrap:wrap; }
.fmt-chip {
    font-size:.56rem; font-weight:700; letter-spacing:.07em; text-transform:uppercase;
    color:var(--text-muted);
    background:rgba(255,255,255,.04); border:1px solid rgba(255,255,255,.07);
    padding:.17rem .48rem; border-radius:2px;
}

/* ── STREAMLIT UPLOADER ── */
[data-testid="stFileUploader"] {
    background: var(--surface) !important;
    border: 1px solid var(--border2) !important;
    border-radius: 4px !important; padding: .25rem !important;
    box-shadow: 0 2px 16px rgba(0,0,0,.4) !important;
    transition: border-color .2s !important;
}
[data-testid="stFileUploader"]:hover {
    border-color: rgba(229,9,20,.35) !important;
}
[data-testid="stFileUploaderDropzone"] {
    background: var(--surface2) !important;
    border: 1.5px dashed rgba(255,255,255,.08) !important;
    border-radius: 3px !important; padding: 1.1rem .75rem !important;
    transition: all .2s !important;
}
[data-testid="stFileUploaderDropzone"]:hover {
    border-color: rgba(229,9,20,.4) !important;
    background: rgba(229,9,20,.025) !important;
}
[data-testid="stFileUploaderDropzoneInstructions"] span {
    color: var(--text-dim) !important; font-size: .78rem !important;
    font-family: 'Barlow', sans-serif !important;
}
[data-testid="stFileUploaderDropzone"] button {
    background: var(--surface3) !important; color: var(--text) !important;
    border: 1px solid rgba(255,255,255,.1) !important; border-radius: 3px !important;
    font-size: .74rem !important; font-weight: 600 !important;
    font-family: 'Barlow', sans-serif !important; padding: .35rem .9rem !important;
    transition: all .15s !important;
}
[data-testid="stFileUploaderDropzone"] button:hover { background: rgba(255,255,255,.1) !important; }
[data-testid="stFileUploader"] label { display: none !important; }

/* ── PRIMARY BUTTON ── */
div.stButton > button {
    width: 100% !important; background: var(--red) !important;
    color: var(--white) !important; font-family: 'Barlow', sans-serif !important;
    font-weight: 700 !important; font-size: .82rem !important;
    letter-spacing: .08em !important; text-transform: uppercase !important;
    border: none !important; border-radius: 3px !important;
    padding: .72rem 1rem !important; margin: 0 !important;
    box-shadow: 0 4px 16px rgba(229,9,20,.3) !important;
    transition: all .18s cubic-bezier(.4,0,.2,1) !important;
}
div.stButton > button:hover {
    background: #f40612 !important;
    box-shadow: 0 6px 28px rgba(229,9,20,.45) !important;
    transform: translateY(-1px) !important;
}
div.stButton > button:active { background: var(--red-dark) !important; }

/* ── ALERTS / SPINNER ── */
.stAlert {
    background: rgba(229,9,20,.08) !important;
    border: 1px solid rgba(229,9,20,.22) !important;
    border-radius: 3px !important; color: #ff7070 !important;
    font-size: .78rem !important; font-weight: 500 !important;
    margin: 0 !important;
}
.stSpinner > div {
    border-color: var(--surface3) var(--surface3) var(--surface3) var(--red) !important;
}

/* ══════════════════════════════════════════════════
   RIGHT PANEL — TABS + DOCUMENT
══════════════════════════════════════════════════ */

/* Tabs header bar */
.stTabs [data-baseweb="tab-list"] {
    background: rgba(0,0,0,.45) !important;
    border-bottom: 1px solid var(--border2) !important;
    gap: 0 !important; padding: 0 1.5rem !important;
    position: sticky; top: 0; z-index: 10;
    margin: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    font-family: 'Barlow', sans-serif !important;
    font-size: .7rem !important; font-weight: 700 !important;
    letter-spacing: .1em !important; text-transform: uppercase !important;
    color: var(--text-muted) !important; background: transparent !important;
    border: none !important; border-bottom: 2px solid transparent !important;
    padding: .72rem 1.2rem !important; margin-bottom: -1px !important;
    transition: color .18s !important;
}
.stTabs [data-baseweb="tab"]:hover { color: var(--text) !important; }
.stTabs [aria-selected="true"] {
    color: var(--white) !important;
    border-bottom: 2px solid var(--red) !important;
    background: transparent !important;
}
.stTabs [data-baseweb="tab-panel"] { padding: 0 !important; }
.stTabs [data-baseweb="tab-highlight"] { display: none !important; }

/* Result header block */
.result-header-block {
    padding: 1rem 1.5rem .6rem;
    display: flex; flex-direction: column; gap: .5rem;
    border-bottom: 1px solid var(--border);
}
.result-header-row {
    display:flex; align-items:center; justify-content:space-between;
}
.result-title {
    font-family:'Bebas Neue', sans-serif; font-size:1.3rem;
    color:var(--white); letter-spacing:.04em;
}
.result-dot {
    width:8px; height:8px; border-radius:50%; background:var(--red);
    animation: pulsedot 2.5s ease infinite; flex-shrink:0;
}
.export-label {
    font-size:.57rem; font-weight:700; letter-spacing:.14em;
    text-transform:uppercase; color:var(--text-muted);
}

/* Download button row */
.dl-button-row {
    display: grid; grid-template-columns: 1fr 1fr 1fr;
    border-bottom: 1px solid var(--border);
}
div.stDownloadButton > button {
    width: 100% !important; background: rgba(255,255,255,.05) !important;
    color: var(--text-dim) !important; font-family: 'Barlow', sans-serif !important;
    font-size: .68rem !important; font-weight: 700 !important;
    letter-spacing: .09em !important; text-transform: uppercase !important;
    border: none !important; border-right: 1px solid var(--border) !important;
    border-radius: 0 !important; padding: .6rem .4rem !important;
    margin: 0 !important; transition: all .15s ease !important;
}
div.stDownloadButton > button:hover {
    background: rgba(255,255,255,.1) !important;
    color: var(--white) !important;
}

/* Document card */
.doc-card-wrap { padding: 1rem 1.5rem 1.5rem; }
.document-card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-top: 2px solid rgba(229,9,20,.25);
    border-radius: 3px;
    padding: 1.25rem 1.5rem;
    max-height: calc(100vh - 260px);
    overflow-y: auto;
    box-shadow: 0 4px 24px rgba(0,0,0,.5);
    animation: fadeUp .35s cubic-bezier(.22,1,.36,1) both;
}
.document-card::-webkit-scrollbar { width: 3px; }
.document-card::-webkit-scrollbar-thumb { background: var(--surface3); border-radius: 3px; }
.document-card pre {
    font-size: .82rem; line-height: 1.8;
    font-family: 'Barlow', sans-serif !important;
    color: rgba(255,255,255,.68);
    white-space: pre-wrap; word-break: break-word;
    margin: 0; padding: 0;
    background: transparent; border: none;
}

/* ── EMPTY STATE ── */
.empty-state {
    display:flex; flex-direction:column;
    align-items:center; justify-content:center;
    text-align:center; padding:3rem;
    height: calc(100vh - var(--chrome-h));
    animation: fadeIn .5s ease both;
}
.empty-icon {
    font-family:'Bebas Neue', sans-serif;
    font-size: clamp(3.5rem, 8vw, 7rem);
    line-height:1; color:rgba(255,255,255,.025);
    letter-spacing:.05em; margin-bottom:1rem;
}
.empty-title { font-size:.85rem; font-weight:600; color:var(--text-muted); }
.empty-sub { font-size:.72rem; color:var(--text-muted); opacity:.5; margin-top:.3rem; }

/* ── MOBILE ── */
@media (max-width: 768px) {
    .nf-date { display: none; }
    [data-testid="stHorizontalBlock"]:first-of-type {
        flex-direction: column !important; height: auto !important;
    }
    [data-testid="stHorizontalBlock"]:first-of-type > [data-testid="column"] {
        width: 100% !important; height: auto !important; overflow: visible !important;
    }
    [data-testid="stHorizontalBlock"]:first-of-type > [data-testid="column"]:first-child {
        border-right: none !important; border-bottom: 1px solid var(--border) !important;
    }
    .document-card { max-height: 55vh; }
    .empty-state { height: 50vh; }
}
</style>
"""

st.markdown(CSS, unsafe_allow_html=True)

# session state
if "detailed" not in st.session_state: st.session_state.detailed = None
if "concise"  not in st.session_state: st.session_state.concise  = None

# ticker
items = ["Meeting Minutes","Flash Reports","PDF Export","Word Export",
         "Multi-File Synthesis","Action Items","Key Decisions","Audio Transcription"]
ticker_html = "".join(
    f'<span class="nf-ticker-item"><b>●</b>{t}</span>' for t in items * 3
)

# ── NAVBAR ────────────────────────────────────────────────────
st.markdown(f"""
<div class="nf-nav">
  <div class="nf-logo">MINUTES<span>AI</span></div>
  <div class="nf-nav-right">
    <span class="nf-date">{current_date_val}</span>
    <span class="nf-pill">Pro</span>
  </div>
</div>
<div class="nf-ticker">
  <div class="nf-ticker-track">{ticker_html}</div>
</div>
""", unsafe_allow_html=True)

# ── TWO-COLUMN LAYOUT via Streamlit columns ───────────────────
left_col, right_col = st.columns([35, 65], gap="small")

# ══ LEFT ══════════════════════════════════════════════════════
with left_col:
    st.markdown("""
    <div class="left-panel">
      <div class="panel-eyebrow">Upload</div>
      <div class="panel-title">MEETING<em> MINUTES</em> ON DEMAND.</div>
      <p class="panel-desc">Drop in audio, transcripts or notes. Walk away with polished professional documents instantly.</p>
      <div class="fmt-row">
        <span class="fmt-chip">MP3</span><span class="fmt-chip">WAV</span>
        <span class="fmt-chip">M4A</span><span class="fmt-chip">TXT</span>
        <span class="fmt-chip">PDF</span><span class="fmt-chip">DOCX</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "upload",
        type=['mp3','wav','m4a','txt','pdf','docx'],
        accept_multiple_files=True,
        label_visibility="collapsed"
    )

    if uploaded_files and st.button("▶  Generate Report"):
        all_context = []
        for f in uploaded_files:
            with st.spinner(f"Processing {f.name}..."):
                if f.name.endswith('.txt'):
                    all_context.append(f.read().decode("utf-8"))
                elif f.name.endswith('.pdf'):
                    all_context.append(extract_text_from_pdf(f))
                elif f.name.endswith('.docx'):
                    all_context.append(extract_text_from_docx(f))
                else:
                    all_context.append(process_audio(f))
        with st.spinner("Synthesizing..."):
            date_instr = f"MANDATORY: Today is {current_date_val}. Replace all [Current Date] placeholders."
            res_det = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=[date_instr, PROFESSIONAL_PROMPT, "\n\n".join(all_context)]
            )
            st.session_state.detailed = res_det.text
            res_con = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=[date_instr, CONCISE_PROMPT, res_det.text]
            )
            st.session_state.concise = res_con.text
        st.success("Done.")

# ══ RIGHT ══════════════════════════════════════════════════════
with right_col:
    if st.session_state.detailed:
        t1, t2 = st.tabs(["  Detailed Minutes  ", "  Executive Flash Report  "])
        for tab, content, title, key_p in zip(
            [t1, t2],
            [st.session_state.detailed, st.session_state.concise],
            ["Detailed Minutes", "Flash Report"],
            ["d", "c"]
        ):
            with tab:
                # Strip markdown for display only; keep original for exports
                display_text = strip_markdown(_html.escape(content))
                # Header block
                st.markdown(f"""
                <div class="result-header-block">
                  <div class="result-header-row">
                    <span class="result-title">{title.upper()}</span>
                    <div class="result-dot"></div>
                  </div>
                  <span class="export-label">Export as</span>
                </div>
                """, unsafe_allow_html=True)
                # Download buttons — tight 3-col row
                c1, c2, c3 = st.columns(3, gap="small")
                c1.download_button("Download Word", create_docx(content, title), f"{title}.docx", key=f"{key_p}_w")
                c2.download_button("Download PDF",  create_pdf(content, title),  f"{title}.pdf",  key=f"{key_p}_p")
                c3.download_button("Download Text", content.encode('utf-8'),     f"{title}.txt",  key=f"{key_p}_t")
                # Document card
                st.markdown(f'<div class="doc-card-wrap"><div class="document-card"><pre>{display_text}</pre></div></div>', unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="empty-state">
          <div class="empty-icon">MINUTES</div>
          <div class="empty-title">No report generated yet</div>
          <div class="empty-sub">Upload files on the left and hit Generate</div>
        </div>
        """, unsafe_allow_html=True)